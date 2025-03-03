import os
import json
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from PIL import Image
from docx.shared import Pt
import pandas as pd
import matplotlib.pyplot as plt

# Constants file path
CONFIG_PATH = "monitoring/config/constants.json"


def load_constants():
    with open(CONFIG_PATH, 'r') as file:
        return json.load(file)


CONSTANTS = load_constants()

def set_document_theme(doc):
    """
    Applies a custom theme to a Word document by setting styles.

    :param doc: The Word Document object.
    """

    normal_style = doc.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Cambria"
    normal_font.size = Pt(11)
    normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")  # Ensures consistency

    # 📌 Explicitly set each heading style
    if "Heading 1" in doc.styles:
        heading1 = doc.styles["Heading 1"]
    else:
        heading1 = doc.styles.add_style("Heading 1", 1)
    heading1_font = heading1.font
    heading1_font.name = "Cambria"
    heading1_font.size = Pt(16)
    heading1_font.bold = True
    heading1.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")

    heading2 = doc.styles["Heading 2"]
    heading2_font = heading2.font
    heading2_font.name = "Cambria"
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    heading2.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")

    heading3 = doc.styles["Heading 3"]
    heading3_font = heading3.font
    heading3_font.name = "Cambria"
    heading3_font.size = Pt(13)
    heading3_font.bold = True
    heading3.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")

    heading4 = doc.styles["Heading 4"]
    heading4_font = heading4.font
    heading4_font.name = "Cambria"
    heading4_font.size = Pt(12)
    heading4_font.italic = True
    heading4.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")

    heading5 = doc.styles["Heading 5"]
    heading5_font = heading5.font
    heading5_font.name = "Cambria"
    heading5_font.size = Pt(11)
    heading5_font.italic = True
    heading5.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")

    # 📌 Check if 'TOC Heading' already exists before adding
    if "TOC Heading" in doc.styles:
        toc_heading = doc.styles["TOC Heading"]
    else:
        toc_heading = doc.styles.add_style("TOC Heading", 1)

    toc_heading_font = toc_heading.font
    toc_heading_font.name = "Cambria"
    toc_heading_font.size = Pt(14)
    toc_heading_font.bold = True
    toc_heading.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")

    # 📌 Customize List Styles
    bullet_list = doc.styles["List Bullet"]
    bullet_list_font = bullet_list.font
    bullet_list_font.name = "Cambria"
    bullet_list_font.size = Pt(11)
    bullet_list.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")

    # 📌 Customize Table Styles
    table_style = doc.styles["Table Grid"]
    table_font = table_style.font
    table_font.name = "Cambria"
    table_font.size = Pt(10)
    table_style.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")

    # 📌 Customize Caption Style (for Figures and Tables)
    caption_style = doc.styles["Caption"]
    caption_font = caption_style.font
    caption_font.name = "Cambria"
    caption_font.size = Pt(10)
    caption_font.bold = True
    caption_style.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")

    # 📌 Customize Footer Style (for Page Numbers)
    footer_style = doc.styles["Footer"]
    footer_font = footer_style.font
    footer_font.name = "Cambria"
    footer_font.size = Pt(10)
    footer_style.element.rPr.rFonts.set(qn('w:eastAsia'), "Cambria")



def add_page_number(doc):
    """Adds page numbers to the footer of the document."""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def add_title_page(doc, report_frequency):
    """Adds the title page."""
    doc.add_paragraph(f"{report_frequency.capitalize()} Monitoring Report", "Title")
    doc.add_page_break()

# def add_doc_control(doc, report_frequency):
#     doc.add_paragraph()


def add_table_of_contents(doc):
    """Adds a TOC field that updates when `F9` is pressed in Word."""
    doc.add_paragraph("Contents", "TOC Heading")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "TOC \\o \"1-3\" \\h \\z \\u"

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    add_list_of_tables_and_figures(doc)

def add_list_of_tables_and_figures(doc):
    """Adds separate 'List of Tables' and 'List of Figures' sections to the document."""

    # 📌 List of Tables
    doc.add_page_break()
    doc.add_paragraph("Tables", "TOC Heading")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "TOC \\h \\z \\t \"Heading 4,1\""  # ✅ Extracts only "Heading 4" for Tables

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    # 📌 List of Figures
    doc.add_paragraph("")
    doc.add_paragraph("Figures", "TOC Heading")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()

    fldChar1_fig = OxmlElement("w:fldChar")
    fldChar1_fig.set(qn("w:fldCharType"), "begin")

    instrText_fig = OxmlElement("w:instrText")
    instrText_fig.set(qn("xml:space"), "preserve")
    instrText_fig.text = "TOC \\h \\z \\t \"Heading 5,1\""  # ✅ Extracts only "Heading 5" for Figures

    fldChar2_fig = OxmlElement("w:fldChar")
    fldChar2_fig.set(qn("w:fldCharType"), "separate")

    fldChar3_fig = OxmlElement("w:fldChar")
    fldChar3_fig.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1_fig)
    run._r.append(instrText_fig)
    run._r.append(fldChar2_fig)
    run._r.append(fldChar3_fig)


def add_list_of_tables(doc):
    """Extracts tables from the document using 'Heading 4' and adds a 'List of Tables' section."""
    doc.add_page_break()
    doc.add_paragraph("Tables", "Title")

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading 4") and "Table" in para.text:
            doc.add_paragraph(para.text, style="Normal")


def add_list_of_figures(doc):
    """Extracts figures from the document using 'Heading 4' and adds a 'List of Figures' section."""
    doc.add_page_break()
    doc.add_paragraph("Figures", "Title")

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading 4") and "Figure" in para.text:
            doc.add_paragraph(para.text, style="Normal")



def replace_placeholders(text, placeholders):
    """Replaces placeholders in text dynamically."""
    if not text:
        return ""

    for key, value in placeholders.items():
        text = text.replace(f"{{{key}}}", str(value))

    return text


def add_section(doc, section_title, section_data, section_number, placeholders, numbering_tracker):
    """Recursively adds sections, subsections, and sub-subsections, while managing numbering of tables and figures at the section level."""

    # Determine section heading level
    heading_level = section_number.count(".") + 1

    # Add a page break only for main sections (level 1 heading)
    if heading_level == 1:
        doc.add_page_break()

    # Fetch the correct section title from JSON (fallback to formatted key name if not found)
    json_title = section_data.get("title", section_title.replace("_", " ").title())

    # Add section heading
    doc.add_heading(f"{section_number}. {json_title}", level=heading_level)

    # Precompute table and figure numbers
    computed_table_numbers, computed_figure_numbers = precompute_numbers(section_data, section_number, numbering_tracker)

    # Replace placeholders and add section text
    process_section_text(doc, section_data, placeholders, computed_table_numbers, computed_figure_numbers)

    # Handle special sections: Scope of Work, Regulatory Standards
    filtered_subsections = process_special_sections(section_title, section_data, placeholders, doc)

    # Add bullet lists
    add_bullet_list(doc, section_data, placeholders)

    # Insert tables, images, and graphs
    insert_tables(doc, section_data, placeholders, computed_table_numbers, computed_figure_numbers)
    insert_images_and_graphs(doc, section_data, computed_figure_numbers, placeholders)

    # Recursively process subsections
    for idx, (sub_key, sub_data) in enumerate(filtered_subsections.items(), start=1):
        sub_section_number = f"{section_number}.{idx}"
        add_section(doc, sub_key, sub_data, sub_section_number, placeholders, numbering_tracker)


def precompute_numbers(section_data, section_number, numbering_tracker):
    """Precompute table and figure numbers before replacing placeholders."""
    main_section_number = section_number.split(".")[0]  # Extract main section (e.g., "4" from "4.1.2")

    computed_table_numbers = []
    computed_figure_numbers = []  # Includes images, graphs, and charts

    # ✅ Precompute Table Numbers
    if "table" in section_data:
        numbering_tracker["table"][main_section_number] = numbering_tracker["table"].get(main_section_number, 0) + 1
        computed_table_numbers.append(f"{main_section_number}.{numbering_tracker['table'][main_section_number]}")
    elif "tables" in section_data:
        num_tables = len(section_data["tables"])
        for _ in range(num_tables):
            numbering_tracker["table"][main_section_number] = numbering_tracker["table"].get(main_section_number, 0) + 1
            computed_table_numbers.append(f"{main_section_number}.{numbering_tracker['table'][main_section_number]}")

    # ✅ Precompute Figure Numbers for Images, Graphs, and Charts
    num_figures = 0
    if "image" in section_data:
        num_figures += 1
    if "images" in section_data:
        num_figures += len(section_data["images"])
    if "graph" in section_data:
        num_figures += 1

    # ✅ Ensure Charts Get a Figure Number
    if "table" in section_data and "data" in section_data["table"]:
        pollutants = section_data["table"]["data"][0][2:]  # Skip 'Monitoring Location' & 'Time' columns
        num_figures += len(pollutants)  # Allocate figure numbers for each pollutant's chart

    if "tables" in section_data:
        for tbl in section_data["tables"]:
            if "data" in tbl:
                pollutants = tbl["data"][0][2:]  # Skip 'Monitoring Location' & 'Time' columns
                num_figures += len(pollutants)

    for _ in range(num_figures):
        numbering_tracker["figure"][main_section_number] = numbering_tracker["figure"].get(main_section_number, 0) + 1
        computed_figure_numbers.append(f"{main_section_number}.{numbering_tracker['figure'][main_section_number]}")

    return computed_table_numbers, computed_figure_numbers




def process_section_text(doc, section_data, placeholders, computed_table_numbers, computed_figure_numbers):
    """Replace placeholders and add section text."""
    text = section_data.get("text", "")

    # Replace `{table_number}` placeholders
    while "{table_number}" in text and computed_table_numbers:
        text = text.replace("{table_number}", computed_table_numbers[0], 1)

    # Replace `{figure_number}` placeholders
    while "{figure_number}" in text and computed_figure_numbers:
        text = text.replace("{figure_number}", computed_figure_numbers[0], 1)

    # Add the final processed text to the document
    if text:
        doc.add_paragraph(replace_placeholders(text, placeholders))


def process_special_sections(section_title, section_data, placeholders, doc):
    """Handle special sections like Scope of Work and Regulatory Standards."""
    if section_title.lower() == "scope of work" and placeholders.get("report_parameters"):
        parameter_list = [p.strip() for p in placeholders["report_parameters"].split(",")]
        formatted_parameters = [format_parameter_section(param) for param in parameter_list]
        for param in formatted_parameters:
            doc.add_paragraph(param, style="List Bullet")

    if section_title.lower() == "regulatory standards" and placeholders.get("report_parameters"):
        parameter_list = [p.strip().lower() for p in placeholders["report_parameters"].split(",")]
        filtered_subsections = {
            key: value for key, value in section_data.get("subsections", {}).items()
            if key.lower() in parameter_list
        }
        if not filtered_subsections:
            print(f"⚠ Warning: No matching regulatory standard found for parameters {parameter_list}.")
        return filtered_subsections

    if section_title.lower() == "conclusion" and placeholders.get("report_parameters"):
        parameter_list = [p.strip().lower() for p in placeholders["report_parameters"].split(",")]

        # Load conclusions and verdict from constants.json
        conclusion_texts = CONSTANTS.get("conclusions", {})
        verdict_text = conclusion_texts.get("verdict",
                                            "This analysis revealed that the observed monitoring parameter(s) consistently adhered to the national standards across all monitored locations at the project site.")

        # Append the relevant conclusions based on selected parameters
        conclusion_paragraphs = []
        for param in parameter_list:
            if param in conclusion_texts:
                conclusion_paragraphs.append(conclusion_texts[param])

        # Add conclusions and verdict to the document
        if conclusion_paragraphs:
            doc.add_paragraph("\n\n".join(conclusion_paragraphs))  # Ensures spacing between paragraphs
            doc.add_paragraph(verdict_text)
        else:
            print(f"⚠ Warning: No matching conclusions found for parameters {parameter_list}.")

    return section_data.get("subsections", {})


def add_bullet_list(doc, section_data, placeholders):
    """Add bullet lists to the document if available."""
    bullet_points = section_data.get("bullet_list", [])
    for point in bullet_points:
        formatted_point = replace_placeholders(point, placeholders)
        doc.add_paragraph(formatted_point, style="List Bullet")



def insert_tables(doc, section_data, placeholders, computed_table_numbers, computed_figure_numbers):
    """Insert tables using precomputed table numbers and dynamically inject monitoring data when applicable."""

    # 🔹 Check if section contains a single table or multiple tables
    if "table" in section_data:
        tables = [section_data["table"]]
    elif "tables" in section_data:
        tables = section_data["tables"]
    else:
        return  # No table data present

    # ✅ Known headers for data injection
    air_quality_headers = ["Monitoring Location", "Time", "CO", "O3", "NO2", "SO2", "PM2.5", "PM10"]
    noise_quality_headers = ["Monitoring Location", "Time", "EQ", "Max", "AE", "10", "50", "90"]

    # 🔹 Check if this section needs dynamic data injection
    if "title" in section_data:
        section_title = section_data["title"].lower()

        # ✅ Monitoring Locations (Scope of Work)
        if "monitoring locations" in section_title and "monitoring_locations" in placeholders:
            tables[0]["data"] = placeholders["monitoring_locations"]

        # ✅ Inject Air & Noise Quality Data
        if "table" in section_data and "data" in section_data["table"] and section_data["table"]["data"]:
            header_row = section_data["table"]["data"][0]

            if header_row == air_quality_headers and "air_monitoring_data" in placeholders:
                section_data["table"]["data"] = placeholders["air_monitoring_data"]

            elif header_row == noise_quality_headers and "noise_monitoring_data" in placeholders:
                section_data["table"]["data"] = placeholders["noise_monitoring_data"]

    # 🔹 Ensure Correct Number of Table Numbers Are Available
    if len(computed_table_numbers) < len(tables):
        print(f"⚠ Warning: Mismatch between precomputed table numbers and actual tables in section.")
        return  # Avoid index errors

    # 🔹 Insert All Tables
    for index, table_data in enumerate(tables):
        if not isinstance(table_data, dict) or "data" not in table_data or not table_data["data"]:
            print(f"⚠ Warning: Unexpected table format in section. Skipping.")
            continue

        table_number = computed_table_numbers[index]  # ✅ Use correct precomputed number
        table_title = replace_placeholders(table_data.get("title", "Table"), placeholders).replace("{table_number}", table_number)

        doc.add_heading(table_title, level=4)

        # ✅ Insert Updated Table into Document
        table = doc.add_table(rows=len(table_data["data"]), cols=len(table_data["data"][0]))
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(table_data["data"]):
            for col_idx, cell_data in enumerate(row_data):
                table.cell(row_idx, col_idx).text = replace_placeholders(str(cell_data), placeholders)

        doc.add_paragraph("")

        if table_data["data"][0] == air_quality_headers or table_data["data"][0] == noise_quality_headers:
            insert_charts(doc, section_data, computed_figure_numbers, placeholders)


def insert_images_and_graphs(doc, section_data, computed_figure_numbers, placeholders):
    """Insert multiple images and graphs with descriptions, ensuring they are centered and appear below."""

    if section_data.get("title") == "Scope of Work":
        monitoring_locations_section = section_data.get("subsections", {}).get("monitoring_locations", {})

        if "images" not in monitoring_locations_section:
            monitoring_locations_section["images"] = []  # Ensure the key exists

        # ✅ Insert Monitoring Location Map at the beginning (if provided)
        if placeholders.get("monitoring_location_map") and not any(
                img.get("description") == "Monitoring Location Map showing marked locations." for img in
                monitoring_locations_section["images"]):
            monitoring_locations_section["images"].insert(0, {
                "path": placeholders["monitoring_location_map"],
                "description": "Environmental Monitoring Location Map"
            })

        # ✅ Insert Site Images for each monitoring location (if provided)
        if placeholders.get("monitoring_location_images"):
            for location, image_path in placeholders["monitoring_location_images"].items():
                monitoring_locations_section["images"].append({
                    "path": image_path,
                    "description": f"Location {location}"
                })

    # 🔹 Handle Multiple Images
    if "images" in section_data:
        for image_data in section_data["images"]:
            if not isinstance(image_data, dict) or "path" not in image_data:
                print("⚠ Warning: Image data format incorrect. Skipping.")
                continue

            if not computed_figure_numbers:
                print("⚠ Warning: Not enough figure numbers for images.")
                continue

            figure_number = computed_figure_numbers.pop(0)
            image_path = image_data["path"]
            image_description = image_data.get("description", f"Figure {figure_number} - Image Description")

            if os.path.exists(image_path):
                try:
                    with Image.open(image_path) as img:
                        dpi = img.info.get('dpi')  # Extract DPI metadata
                        if dpi is None or dpi[0] == 0 or dpi[1] == 0:
                            print(f"⚠ Warning: DPI metadata missing for {image_path}. Setting default DPI to 96.")
                            dpi = (96, 96)
                        img.save(image_path, dpi=dpi)  # Save with corrected DPI

                    # 🔹 Determine Image Size
                    image_width = Inches(5) if "Location Map" in image_description else Inches(2.5)  # Larger for Location Map

                    # 🔹 Insert Image and Center Align
                    image_paragraph = doc.add_paragraph()
                    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = image_paragraph.add_run()
                    run.add_picture(image_path, width=image_width)  # ✅ Adjust width dynamically

                    # 🔹 Add Image Description Below
                    desc_paragraph = doc.add_heading(f"Figure {figure_number} - {image_description}", level=5)
                    desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.add_paragraph("")

                except Exception as e:
                    print(f"⚠ Warning: Failed to insert image {image_path}. Error: {e}")

    # 🔹 Handle Single Image (for backward compatibility)
    elif "image" in section_data:
        figure_number = computed_figure_numbers.pop(0)
        image_path = section_data["image"]
        image_description = section_data.get("image_description", f"Figure {figure_number} - Image Description")

        if os.path.exists(image_path):
            try:
                with Image.open(image_path) as img:
                    dpi = img.info.get('dpi')
                    if dpi is None or dpi[0] == 0 or dpi[1] == 0:
                        print(f"⚠ Warning: DPI metadata missing for {image_path}. Setting default DPI to 96.")
                        dpi = (96, 96)
                    img.save(image_path, dpi=dpi)

                # 🔹 Determine Image Size
                image_width = Inches(3) if "Location Map" in image_description else Inches(1.5)  # Larger for Location Map

                # 🔹 Insert Image and Center Align
                image_paragraph = doc.add_paragraph()
                image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = image_paragraph.add_run()
                run.add_picture(image_path, width=image_width)

                # 🔹 Add Image Description Below
                desc_paragraph = doc.add_paragraph(f"Figure {figure_number} - {image_description}")
                desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph("")

            except Exception as e:
                print(f"⚠ Warning: Failed to insert image {image_path}. Error: {e}")

def insert_charts(doc, section_data, computed_figure_numbers, placeholders):
    """Generate and insert charts for air and noise quality monitoring data using sequential figure numbering."""

    # Define headers for air and noise quality
    air_quality_headers = ["Monitoring Location", "Time", "CO", "O3", "NO2", "SO2", "PM2.5", "PM10"]
    noise_quality_headers = ["Monitoring Location", "Time", "EQ", "Max", "AE", "10", "50", "90"]  # Ensure correct case

    # Identify if air or noise monitoring data is present
    if "table" in section_data and "data" in section_data["table"]:
        table_data = section_data["table"]["data"]
    elif "tables" in section_data:
        table_data = section_data["tables"][0]["data"] if section_data["tables"] else []
    else:
        print("⚠ Warning: No relevant table data found.")
        return  # No relevant table data

    # Convert table data to DataFrame
    df = pd.DataFrame(table_data[1:], columns=table_data[0])  # Use first row as headers

    # ✅ Determine monitoring type
    if set(df.columns) == set(air_quality_headers):  # Using set() to ignore column order issues
        monitoring_type = "Air Quality"
        ncec_standards = {
            "CO": 40000,
            "O3": 157,
            "NO2": 200,
            "SO2": 441,
            "PM2.5": 35,
            "PM10": 340
        }
        y_axis_label = "Concentration (μg/m³)"  # ✅ Air quality uses μg/m³
    elif set(df.columns) == set(noise_quality_headers):  # ✅ Ensuring we match the noise headers correctly
        monitoring_type = "Noise Quality"
        ncec_standards = {"EQ": 70}  # ✅ Only EQ benchmark (70 dB)
        y_axis_label = "Noise Level (dB)"  # ✅ Noise quality uses dB
    else:
        print(f"⚠ Warning: Table headers do not match expected Air/Noise quality formats. Headers found: {df.columns.tolist()}")
        return  # Not an air/noise monitoring table

    # ✅ Extract monitoring locations dynamically
    locations = df["Monitoring Location"].tolist()

    # ✅ Extract pollutants dynamically (Only EQ for Noise)
    if monitoring_type == "Air Quality":
        pollutants = [col for col in df.columns if col not in ["Monitoring Location", "Time"]]
    elif monitoring_type == "Noise Quality":
        pollutants = ["EQ"] if "EQ" in df.columns else []  # ✅ Only include EQ for Noise

    # ✅ Ensure we have valid pollutants to plot
    if not pollutants:
        print(f"⚠ Warning: No valid pollutants found for {monitoring_type}. Skipping chart generation.")
        return

    # ✅ Generate and save charts dynamically
    saved_files = []
    for pollutant in pollutants:
        if not computed_figure_numbers:
            print(f"⚠ Warning: Not enough figure numbers for charts.")
            continue

        figure_number = computed_figure_numbers.pop(0)  # Fetch the next figure number

        fig, ax = plt.subplots(figsize=(6, 4))  # Set figure size

        # Extract pollutant values for each location
        pollutant_values = df[pollutant].astype(float).tolist()

        # ✅ Plot bars
        bars = ax.bar(locations, pollutant_values, color='#1f77b4', width=0.4, label=f"{pollutant} Levels")

        # ✅ Add a horizontal benchmark line if applicable
        if pollutant in ncec_standards:
            ax.axhline(y=ncec_standards[pollutant], color='red', linestyle='--', linewidth=2,
                       label=f"NCEC Std. ({ncec_standards[pollutant]} {y_axis_label})")

        # ✅ Labels and title
        ax.set_xlabel("Monitoring Locations")
        ax.set_ylabel(y_axis_label)
        ax.set_title(f"{monitoring_type} - {pollutant} Levels")
        ax.legend()

        # ✅ Save the figure
        filename = f"{pollutant.replace(' ', '_').replace('/', '_')}_levels.png"
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        saved_files.append((filename, figure_number, pollutant))

        plt.close(fig)  # ✅ Prevent display when running script

    # ✅ Insert images into Word document
    for image_path, figure_number, pollutant in saved_files:
        doc.add_heading(f"Figure {figure_number} - {monitoring_type} - {pollutant} Levels", level=5)

        # ✅ Insert Image and Center Align
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = image_paragraph.add_run()
        run.add_picture(image_path, width=Inches(4))

        doc.add_paragraph("")  # ✅ Add spacing below

        # ✅ Clean up temporary files
        os.remove(image_path)




def format_parameter_section(parameter):
    """Formats user input parameters into proper section titles."""
    formatted_parameters = {
        "air": "Ambient Air Quality Monitoring",
        "noise": "Noise Monitoring",
        "soil": "Soil Quality Monitoring",
        "water": "Water Quality Monitoring"
    }
    return formatted_parameters.get(parameter.lower(), parameter.capitalize() + " Monitoring")

def generate_report():
    """Generates a monitoring report dynamically based on input data."""

    structure_file = CONSTANTS["structure_file"]


    placeholders = {'consultancy_name': 'Green Fields Environmental Consulting',
                    'contractor_name': 'Abdullah Bin Talib for Swimming Pools Co.',
                    'project_name': 'Concrete Structure & Civil Works of the Marina Lifestyle Hotel asset',
                    'project_number': 'PR2408074 ',
                    'reference_number': '2408074-RSG-MAC-WR-23',
                    'report_frequency': 'Weekly',
                    'report_date': '05 Jan 2025',
                    'report_number': '59th',
                    'report_parameters': 'Air, Noise',
                    'monitoring_frequency': '30 mins',
                    'monitoring_locations': [['Monitoring Location', 'Description', 'Latitude', 'Longitude'],
                                             ['ML-01', 'Family Pool', '26.636180°', '36.224574°'],
                                             ['ML-02', 'Couple Pool', '26.627794°', '36.227677°']],
                    'monitoring_location_map': 'monitoring/test_data/map.png',
                    'monitoring_location_images': {'ML-01': 'monitoring/test_data/ml01.png',
                                                   'ML-02': 'monitoring/test_data/ml02.png'},
                    'air_monitoring_data': [['Monitoring Location', 'Time', 'CO', 'O3', 'NO2', 'SO2', 'PM2.5', 'PM10'],
                                            ['ML-01', '30/12/2024 09:37', '1016.4', '51', '88.8', '41.4', '14.3', '120.9'],
                                            ['ML-02', '30/12/2024 10:22', '1253.3', '37.0', '64.2', '99.8', '15.8', '131.3']],
                    'noise_monitoring_data': [['Monitoring Location', 'Time', 'EQ', 'Max', 'AE', '10', '50', '90'],
                                              ['ML-01', '30/12/2024 09:37', '61.3', '72.3', '93.9', '64.1', '60.06', '55.8'],
                                              ['ML-02', '30/12/2024 10:22', '61', '82.3', '93.6', '64.2', '58.6', '55.8']]}

    # Load structured JSON
    with open(structure_file, 'r') as file:
        section_data = json.load(file)

    # Convert JSON keys to lowercase for **case-insensitive** lookup
    section_data = {key.lower(): value for key, value in section_data.items()}

    # Define **explicit** section order
    sections = [
        "Introduction",
        "Scope of Work",
        "Regulatory Standards"
    ]

    # 📌 Insert User-defined Sections Dynamically
    if placeholders["report_parameters"]:
        parameter_list = [p.strip() for p in placeholders["report_parameters"].split(",")]
        formatted_parameters = [format_parameter_section(param) for param in parameter_list]
        sections.extend(formatted_parameters)

    # 📌 Add Final Sections
    sections.extend(["Conclusion", "Appendices"])

    # Initialize table/figure/graph numbering tracker
    numbering_tracker = {
        "table": {},
        "figure": {},
        "graph": {}
    }

    # Create Word document
    doc = Document()
    # set_document_theme(doc)


    add_page_number(doc)

    # 📌 Title Page
    # add_title_page(doc, placeholders["report_frequency"])

    # 📌 Table of Contents
    add_table_of_contents(doc)

    # 📌 Generate Sections
    for i, section_key in enumerate(sections, start=1):
        # Convert key to lowercase for safe lookup
        section_key_lower = section_key.lower().replace(" ", "_")

        if section_key_lower in section_data:
            section_title = section_key  # ✅ Fetch correct title
            add_section(doc, section_title, section_data[section_key_lower], str(i), placeholders, numbering_tracker)
        else:
            print(f"⚠ Warning: Section '{section_key}' not found in JSON.")


    # 📌 Save Document
    output_dir = CONSTANTS["output_dir"]
    os.makedirs(output_dir, exist_ok=True)
    report_path = f"{output_dir}/{placeholders['report_frequency'].capitalize()}_Monitoring_Report.docx"
    doc.save(report_path)

    print(f"✅ {placeholders['report_frequency'].capitalize()} Monitoring Report generated: {report_path}")
    return report_path