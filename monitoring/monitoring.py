import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
import os
import json
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# Constants file path
CONFIG_PATH = "monitoring/config/constants.json"

def load_constants():
    with open(CONFIG_PATH, 'r') as file:
        return json.load(file)

CONSTANTS = load_constants()

def add_page_number(doc):
    """Adds page numbers to the footer of the document."""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center alignment
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def add_title_page(doc, report_frequency):
    """Adds the title page."""
    doc.add_paragraph(f"{report_frequency.capitalize()} Monitoring Report", "Title")
    doc.add_page_break()

def add_table_of_contents(doc):
    """Adds a TOC field that updates when `F9` is pressed in Word."""
    doc.add_paragraph("Table of Contents", "Title")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align TOC for better presentation
    run = paragraph.add_run()

    # Word TOC Field Code
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "TOC \\o \"1-3\" \\h \\z \\u"

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def add_section(doc, section_title, section_number):
    """Adds a section."""
    doc.add_page_break()
    doc.add_heading(f"{section_number}. {section_title}", level=1)

def add_subsection(doc, section_title, section_number):
    """Adds a subsection with level 2 heading."""
    doc.add_heading(f"{section_number}. {section_title}", level=2)


def add_subsection_content(doc, data, subsection_key):
    """Adds content dynamically to a subsection, including text, bullet lists, images, tables, and graphs."""

    # Fetch content
    content = data.get(subsection_key, {})

    # Add text (if available)
    text = content.get("text", "")
    if text:
        doc.add_paragraph(text)

    # Add bullet list (if available)
    bullet_points = content.get("bullet_points", [])
    for point in bullet_points:
        doc.add_paragraph(point, style="List Bullet")

    # Add table (if available)
    table_data = content.get("table")
    if table_data:
        table_title = table_data.get("title", "Table")
        doc.add_paragraph(table_title, style="Heading 3")

        table = doc.add_table(rows=len(table_data["data"]), cols=len(table_data["data"][0]))
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(table_data["data"]):
            for col_idx, cell_data in enumerate(row_data):
                table.cell(row_idx, col_idx).text = str(cell_data)

    # Add image (if available)
    image_path = content.get("image")
    if image_path and os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5))

    # Add graph (if available)
    graph_path = content.get("graph")
    if graph_path and os.path.exists(graph_path):
        doc.add_picture(graph_path, width=Inches(5))


def add_introduction(doc, constant_text_json_file, consultancy_name, contractor_name, project_name, report_frequency,
                     report_date, report_number, section_number):
    """Adds the introduction section."""
    add_section(doc, "Introduction", section_number)

    with open(constant_text_json_file, 'r') as file:
        data = json.load(file)
    first_paragraph = data.get("introduction_first_paragraph", "No introduction text available.")
    second_paragraph = data.get("introduction_second_paragraph", "")
    first_paragraph = (first_paragraph
                       .replace("{consultancy_name}", consultancy_name)
                       .replace("{contractor_name}", contractor_name)
                       .replace("{project_name}", project_name)
                       .replace("{report_frequency}", report_frequency.capitalize()))
    second_paragraph = (second_paragraph
                        .replace("{report_date}", report_date)
                        .replace("{report_number}", report_number)
                        .replace("{report_frequency}", report_frequency.capitalize()))
    doc.add_paragraph(first_paragraph)
    doc.add_paragraph(second_paragraph)

def add_scope_of_work(doc, constant_text_json_file, section_number, report_parameters, report_date):
    """Adds the Scope of Work section with static content and a bullet list of parameters."""
    subsections = ["Monitoring Locations"]
    table_count = 0
    figure_count = 0

    add_section(doc, "Scope of Work", section_number)

    # Load scope of work text from JSON
    with open(constant_text_json_file, 'r') as file:
        data = json.load(file)
    scope_text = data.get("scope_of_work", "No scope of work text available.")

    doc.add_paragraph(scope_text)

    # Format parameters into proper section names
    if report_parameters:

        parameter_list = [p.strip() for p in report_parameters.split(",")]
        formatted_parameters = [format_parameter_section(param) for param in parameter_list]

        for param in formatted_parameters:
            doc.add_paragraph(f"{param}", style="List Bullet")

    for idx, subsection in enumerate(subsections, start=1):
        subsection_number = f"{section_number}.{idx}"
        add_subsection(doc, subsection, subsection_number)

        # Fetch subsection content from JSON
        subsection_key = f"scope_of_work_{subsection.lower().replace(' ', '_')}"
        add_subsection_content(doc, data, subsection_key)


def add_regulatory_standards(doc, constant_text_json_file, section_number, report_parameters):
    """Adds the Regulatory Standards section with dynamic subsections based on report parameters."""

    table_count = 0
    figure_count = 0

    add_section(doc, "Regulatory Standards", section_number)

    # Load regulatory standards text from JSON
    with open(constant_text_json_file, 'r') as file:
        data = json.load(file)

    # 📌 Process report parameters dynamically
    parameter_list = [p.strip().lower() for p in report_parameters.split(",") if p.strip()]

    for idx, param in enumerate(parameter_list, start=1):
        formatted_param = format_parameter_section(param)  # Use existing function
        subsection_title = f"Regulatory Standard - {formatted_param}"

        add_subsection(doc, subsection_title, f"{section_number}.{idx}")

        # Load content for the subsection from JSON
        subsection_key = f"regulatory_standard_{param}"
        subsection_text = data.get(subsection_key, "No specific regulatory details available.")

        doc.add_paragraph(subsection_text)


def add_air_quality_monitoring(doc, introduction_json_file, section_number):
    """Adds the Ambient Air Quality Monitoring section with dynamically generated subsections."""

    subsections = [
        "Objective",
        "Scope",
        "Instrumentation and Methodology",
        "Results and Discussions"
    ]

    add_section(doc, "Ambient Air Quality Monitoring", section_number)

    # Load air quality monitoring content from JSON
    with open(introduction_json_file, 'r') as file:
        data = json.load(file)

    # General introduction to the section
    air_quality_text = data.get("air_quality_monitoring", "No ambient air quality monitoring data available.")
    doc.add_paragraph(air_quality_text)


    # Loop through and add subsections
    for idx, subsection in enumerate(subsections, start=1):
        subsection_number = f"{section_number}.{idx}"
        add_subsection(doc, subsection, subsection_number)

        # Fetch subsection content from JSON
        subsection_key = f"air_quality_{subsection.lower().replace(' ', '_')}"
        add_subsection_content(doc, data, subsection_key)


def format_parameter_section(parameter):
    """Formats user input parameters into proper section titles."""
    formatted_parameters = {
        "air": "Ambient Air Quality",
        "noise": "Noise",
        "soil": "Soil Quality",
        "water": "Water Quality"
    }
    return formatted_parameters.get(parameter.lower(), parameter.capitalize() + " Monitoring")

def generate_report():
    """Generates a monitoring report based on user input."""
    contractor_name = input("Enter the contractor's name: ").strip()
    project_name = input("Enter the project name: ").strip()
    report_frequency = input("Enter report frequency (Weekly, Monthly): ").strip().lower()
    report_date = input("Enter the report date (e.g., 06th January 2025): ").strip()
    report_number = input("Enter the report number (e.g., Twenty-third): ").strip()
    report_parameters = input("Enter report parameters (comma-separated, e.g., Air, Noise, Soil): ").strip()


    # Load introduction content from JSON
    constant_text_json_file = CONSTANTS["constant_text_json_file"]
    consultancy_name = CONSTANTS["consultancy_name"]

    # Create a new Word document
    doc = Document()
    add_page_number(doc)

    # 📌 1. Title Page (Page 1)
    add_title_page(doc, report_frequency)

    # 📌 2. Table of Contents (Page 2)
    add_table_of_contents(doc)

    # 📌 3. Build Sections List
    sections = [
        "Introduction",
        "Scope of Work",
        "Regulatory Standards"
    ]

    # 📌 4. Insert User-defined Sections Dynamically
    if report_parameters:
        parameter_list = [p.strip() for p in report_parameters.split(",")]
        formatted_parameters = [format_parameter_section(param) for param in parameter_list]
        sections.extend(formatted_parameters)

    # 📌 5. Add Final Sections
    sections.extend(["Conclusion", "Appendices"])

    # 📌 6. Generate Sections
    for i, section in enumerate(sections, start=1):
        if section == "Introduction":
            add_introduction(doc, constant_text_json_file, consultancy_name, contractor_name, project_name,
                             report_frequency, report_date, report_number, i)
        elif section == "Scope of Work":
            add_scope_of_work(doc, constant_text_json_file, i, report_parameters, report_date)
        elif section == "Regulatory Standards":
            add_regulatory_standards(doc, constant_text_json_file, i, report_parameters)
        elif section == "Ambient Air Quality Monitoring":
            add_air_quality_monitoring(doc, constant_text_json_file, i)


        else:
            add_section(doc, section, i)

    # Save the document
    output_dir = CONSTANTS["output_dir"]
    os.makedirs(output_dir, exist_ok=True)
    report_path = f"{output_dir}/{report_frequency.capitalize()}_Monitoring_Report.docx"
    doc.save(report_path)

    print(f"{report_frequency.capitalize()} Monitoring Report generated: {report_path}")


